import docx
from docx.enum.text import WD_COLOR_INDEX
import os
import shutil
import glob

appendix_dir = "/Users/congnguyen/Library/CloudStorage/OneDrive-Personal/Work/04. Company/Synodus/GOV/01.VNA/Document/Phụ lục"
backup_dir = os.path.join(appendix_dir, "backup_original")

# 1. Restore from backup to start with clean original documents
backup_files = glob.glob(os.path.join(backup_dir, "*.docx"))
print(f"Restoring {len(backup_files)} original files from backup...")
for f in backup_files:
    fname = os.path.basename(f)
    dest_path = os.path.join(appendix_dir, fname)
    shutil.copy2(f, dest_path)
print("Restore complete. Starting modification with highlight...\n")

# 2. Define the replacement data (same approved text as before)
DEPT_DATA = {
    "KHPT": {
        "title": "Tính năng Dashboard tổng quan, Quản lý mục tiêu KPI và Kho tài liệu PTBV chung.",
        "scope": "Các tính năng hiển thị Dashboard tổng quan chung, 3 trụ cột E, S, G, Quản lý KPI và Kho tài liệu PTBV chung của Ban KHPT.",
        "requirements": [
            ['YCQ-KHPT-WEB-01', 'Giao diện Dashboard điều hành tổng quan chung (Executive Dashboard).', 'Cao', 
             'Hiển thị biểu đồ phân tích tích hợp (Metabase iframe) gồm 4 tab chuyển đổi: Tổng quan chung (ALL), Environmental (E), Social (S), và Governance (G). Hỗ trợ nút xem chi tiết trên Metabase và bảng danh sách chỉ tiêu chi tiết bên dưới.'],
            ['YCQ-KHPT-WEB-02', 'Quy trình Đề xuất & Thiết lập mục tiêu KPI.', 'Cao', 
             'Giao diện hỗ trợ quản lý danh mục KPI có các trường: Mã chỉ tiêu, Tên chỉ tiêu, Đơn vị tính, Số kế hoạch, Số thực hiện, Tiến độ (%), Đánh giá (Đạt/Không đạt), và Đơn vị chủ trì. Cho phép KHPT gửi yêu cầu làm kế hoạch KPI năm xuống các phòng ban, theo dõi tiến độ nộp và phê duyệt/yêu cầu nộp lại số liệu kế hoạch trực tiếp trên hệ thống (Kế hoạch lưu vào mục riêng cho KHPT quản lý).'],
            ['YCQ-KHPT-WEB-03', 'Quản trị Kho tài liệu Phát triển bền vững chung.', 'Trung bình', 
             'Tiếp nhận yêu cầu upload tài liệu của các phòng ban dưới dạng tệp đính kèm. Hỗ trợ KHPT duyệt lưu trữ (vào thư mục dùng chung) hoặc từ chối chỉnh sửa.']
        ],
        "traceability": [
            ['YCQ-KHPT-WEB-01', 'Hình 01: Giao diện dashboard tổng quan', 'Dữ liệu Metabase nhúng', 'Kiểm thử hiển thị biểu đồ và bảng danh sách chỉ tiêu.'],
            ['YCQ-KHPT-WEB-02', 'Hình 02: Màn hình quản lý KPI', 'Dữ liệu chỉ tiêu KPI', 'Kiểm thử cập nhật số kế hoạch và theo dõi tiến độ.'],
            ['YCQ-KHPT-WEB-03', 'Hình 03: Màn hình kho tài liệu', 'Tệp đính kèm pdf/docx', 'Kiểm thử tải lên tài liệu và nút kiểm duyệt lưu trữ.']
        ],
        "uat": [
            ['YCQ-KHPT-WEB-01', 'Kiểm thử Dashboard tổng quan', 'Mở trang Tổng quan điều hành, kiểm tra hiển thị iframe Metabase và các tab E, S, G. Kiểm tra bảng chỉ tiêu bên dưới.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-KHPT-WEB-02', 'Kiểm thử quản lý KPI', 'Mở chức năng Quản lý KPI, kiểm tra nhập số kế hoạch KPI cho các chỉ tiêu.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-KHPT-WEB-03', 'Kiểm thử Kho tài liệu', 'Mở chức năng Kho tài liệu PTBV chung, kiểm tra tải lên tài liệu và kiểm duyệt lưu trữ.', 'Đạt', '—', 'Hoàn thành']
        ]
    },
    "Truyền thông": {
        "title": "Tính năng CMS quản trị Web Public, API Spirit VNA và Lưu trữ báo cáo ESG.",
        "scope": "Các tính năng hiển thị CMS Website ESG, đồng bộ API Spirit VNA, Quản lý báo cáo ESG tập trung và chỉ tiêu GRI 417-3.",
        "requirements": [
            ['YCQ-BTT-WEB-01', 'CMS Quản trị Web Public ESG.', 'Cao', 'Cho phép Ban TT chỉnh sửa, biên soạn và duyệt xuất bản các bài viết Chiến lược ESG (3 trụ cột), Thông điệp/Tầm nhìn của ban lãnh đạo. Hỗ trợ cơ chế ẩn/hiện bài viết trên trang Web Public.'],
            ['YCQ-BTT-WEB-02', 'API đồng bộ tin tức hoạt động.', 'Cao', 'Tích hợp API tự động đồng bộ các bài viết từ cổng tin tức nội bộ Spirit VNA về danh mục Tin tức hoạt động của Website ESG Public.'],
            ['YCQ-BTT-WEB-03', 'Lưu trữ báo cáo ESG tập trung.', 'Trung bình', 'Cho phép Ban TT gửi yêu cầu nộp báo cáo thường niên ESG tới các phòng ban. Hệ thống hiển thị danh mục các báo cáo được nộp tập trung dưới dạng thư mục, hỗ trợ tải về trực tiếp.'],
            ['YCQ-BTT-WEB-04', 'Hiển thị tuyên bố mặc định GRI 417-3.', 'Cao', 'Khi số vụ vi phạm bằng 0, hệ thống tự động hiển thị câu tuyên bố: \"Vietnam Airlines không có vụ việc nào liên quan đến tuân thủ truyền thông marketing.\" thay vì để trống dữ liệu.']
        ],
        "traceability": [
            ['YCQ-BTT-WEB-01', 'Hình 01: Màn hình quản trị CMS', 'Nội dung bài viết ESG', 'Kiểm thử soạn thảo, biên tập và duyệt xuất bản bài viết.'],
            ['YCQ-BTT-WEB-02', 'Hình 02: Trang tin tức hoạt động', 'API Feed từ Spirit VNA', 'Kiểm thử đồng bộ dữ liệu tự động.'],
            ['YCQ-BTT-WEB-03', 'Hình 03: Thư mục báo cáo ESG', 'Tệp PDF báo cáo', 'Kiểm thử yêu cầu nộp và tải xuống báo cáo.'],
            ['YCQ-BTT-WEB-04', 'Hình 04: Biểu đồ báo cáo GRI 417-3', 'Chỉ số vi phạm = 0', 'Kiểm thử hiển thị câu tuyên bố mặc định.']
        ],
        "uat": [
            ['YCQ-BTT-WEB-01', 'Kiểm thử CMS quản trị', 'Thực hiện tạo bài viết mới, biên tập và xuất bản lên trang public. Kiểm tra nút ẩn/hiện bài viết.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-BTT-WEB-02', 'Kiểm thử đồng bộ API', 'Kiểm tra danh mục tin tức tự động cập nhật bài viết mới từ Spirit VNA.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-BTT-WEB-03', 'Kiểm thử Lưu trữ báo cáo', 'Gửi yêu cầu nộp báo cáo và kiểm tra tệp báo cáo thường niên xuất hiện tại thư mục tập trung.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-BTT-WEB-04', 'Kiểm thử Hiển thị chỉ tiêu', 'Kiểm tra giao diện hiển thị câu tuyên bố mặc định của GRI 417-3 khi số vụ vi phạm bằng 0.', 'Đạt', '—', 'Hoàn thành']
        ]
    },
    "Dịch vụ": {
        "title": "Tính năng nhập liệu Dịch vụ khách hàng, NPS và sự cố dịch vụ & nhãn mác.",
        "scope": "Các tính năng nhập mục tiêu năm, định mức nước, nhật ký sự cố dịch vụ & nhãn mác, và Dashboard NPS của Tổ Dịch vụ.",
        "requirements": [
            ['YCQ-TDV-WEB-01', 'Nhập liệu mục tiêu năm.', 'Cao', 'Form nhập liệu các chỉ tiêu nước sạch năm (waterWithdrawalTarget, waterConsumptionTarget), Tổng số nhà cung cấp (totalSuppliers), Số nhà cung cấp địa phương (localSuppliers), Điểm NPS kế hoạch (npsTarget).'],
            ['YCQ-TDV-WEB-02', 'Nhật ký sự cố vi phạm an toàn dịch vụ (GRI 416-2) và vi phạm nhãn mác (GRI 417-2).', 'Cao', 'Giao diện cho phép cập nhật danh sách sự cố gồm: Mã sự cố, Tháng xảy ra, Mô tả vụ việc. Nếu danh sách trống, hiển thị tuyên bố mặc định: \"Không có sự cố nào trong kỳ báo cáo.\"'],
            ['YCQ-TDV-WEB-03', 'Dashboard NPS & Nguồn nước.', 'Cao', 'Trực quan hóa dữ liệu NPS và lượng nước tiêu thụ/tuần hoàn theo tháng/quý/năm trên dashboard chuyên môn.']
        ],
        "traceability": [
            ['YCQ-TDV-WEB-01', 'Hình 01: Màn hình Mục tiêu năm', 'Dữ liệu mục tiêu năm', 'Kiểm thử nhập và lưu các chỉ số mục tiêu năm.'],
            ['YCQ-TDV-WEB-02', 'Hình 02: Màn hình Nhật ký sự cố', 'Danh sách sự cố', 'Kiểm thử nhập danh sách sự cố và hiển thị tuyên bố mặc định khi sự cố bằng 0.'],
            ['YCQ-TDV-WEB-03', 'Hình 03: Dashboard Dịch vụ & NPS', 'Biểu đồ Metabase', 'Kiểm thử hiển thị biểu đồ xu hướng NPS và tiêu thụ nước.']
        ],
        "uat": [
            ['YCQ-TDV-WEB-01', 'Kiểm thử nhập mục tiêu năm', 'Nhập các giá trị mục tiêu nước và NPS kế hoạch, kiểm tra lưu trữ thành công.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-TDV-WEB-02', 'Kiểm thử nhật ký sự cố', 'Nhập một vụ vi phạm nhãn mác mới, kiểm tra hiển thị. Xóa hết sự cố để kiểm tra câu tuyên bố mặc định.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-TDV-WEB-03', 'Kiểm thử Dashboard Dịch vụ', 'Kiểm tra hiển thị các biểu đồ NPS nội địa/quốc tế và tỷ lệ Promoter.', 'Đạt', '—', 'Hoàn thành']
        ]
    },
    "CĐSCN": {
        "title": "Tính năng nhập liệu sự cố an toàn thông tin và bảo mật dữ liệu.",
        "scope": "Các tính năng nhập liệu chỉ tiêu GRI 418-1 của Ban CĐSCN và cảnh báo vi phạm bảo mật dữ liệu.",
        "requirements": [
            ['YCQ-CDS-WEB-01', 'Nhập liệu sự cố dữ liệu cá nhân (GRI 418-1).', 'Cao', 'Form nhập liệu thủ công hàng tháng đối với 2 tham số: Số khiếu nại vi phạm quyền riêng tư (privacyBreaches) và Số vụ mất mát dữ liệu cá nhân (dataLosses). Có tính năng kiểm tra lỗi định dạng dữ liệu đầu vào.']
        ],
        "traceability": [
            ['YCQ-CDS-WEB-01', 'Hình 01: Màn hình Nhập liệu bảo mật', 'Số liệu sự cố bảo mật', 'Kiểm thử nhập số vụ vi phạm và kiểm tra định dạng dữ liệu.']
        ],
        "uat": [
            ['YCQ-CDS-WEB-01', 'Kiểm thử nhập số liệu bảo mật', 'Nhập các số vụ vi phạm bảo mật dữ liệu và mất mát thông tin, kiểm tra lưu dữ liệu và hiển thị cảnh báo.', 'Đạt', '—', 'Hoàn thành']
        ]
    },
    "Khai thác": {
        "title": "Tính năng nhập liệu Khai thác bay, mức ồn, phát thải và nhật ký sự cố.",
        "scope": "Các tính năng nhập tiêu hao nhiên liệu, mức ồn, phát thải khí nhà kính và nhật ký sự cố khai thác bay của Tổ Khai thác (không hiển thị chỉ tiêu khí NOx, SOx GRI 305-7).",
        "requirements": [
            ['YCQ-TKT-WEB-01', 'Nhật ký sự cố an toàn khai thác bay.', 'Cao', 'Nhập liệu và quản lý danh sách sự cố an toàn khai thác bay bao gồm: Mã sự cố, Ngày xảy ra, Tiêu đề sự cố, Trạng thái xử lý (Open | Investigating | Closed).'],
            ['YCQ-TKT-WEB-02', 'Rút chỉ tiêu khí thải NOx, SOx (GRI 305-7).', 'Cao', 'Loại bỏ hoàn toàn chỉ tiêu phát thải khí NOx, SOx (GRI 305-7) ra khỏi giao diện báo cáo và dashboard khai thác trên Web Portal.'],
            ['YCQ-TKT-WEB-03', 'Dashboard năng lượng & phát thải Scope 1.', 'Cao', 'Trực quan hóa lượng Jet Fuel & SAF tiêu thụ, lượng CO2 phát thải hóa thạch và biogenic CO2. Hỗ trợ bộ lọc chặng bay, loại tàu bay, và khu vực bay.']
        ],
        "traceability": [
            ['YCQ-TKT-WEB-01', 'Hình 01: Màn hình Nhật ký sự cố', 'Danh sách sự cố khai thác', 'Kiểm thử nhập liệu sự cố và cập nhật trạng thái xử lý.'],
            ['YCQ-TKT-WEB-02', 'Hình 02: Danh mục chỉ tiêu Khai thác', 'Danh sách chỉ tiêu', 'Xác nhận không còn xuất hiện mã chỉ tiêu GRI 305-7.'],
            ['YCQ-TKT-WEB-03', 'Hình 03: Dashboard Khai thác bay', 'Biểu đồ Metabase', 'Kiểm thử hiển thị biểu đồ tiêu hao nhiên liệu và phát thải CO2.']
        ],
        "uat": [
            ['YCQ-TKT-WEB-01', 'Kiểm thử nhật ký sự cố', 'Nhập sự cố va chạm chim bay, kiểm tra lưu trữ và thay đổi trạng thái sang Closed.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-TKT-WEB-02', 'Kiểm thử rút chỉ tiêu SOx/NOx', 'Kiểm tra màn hình chỉ tiêu và dashboard để đảm bảo chỉ tiêu GRI 305-7 đã được ẩn hoàn toàn.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-TKT-WEB-03', 'Kiểm thử Dashboard Khai thác', 'Kiểm tra hiển thị dữ liệu lượng phát thải CO2 Scope 1 Domestic.', 'Đạt', '—', 'Hoàn thành']
        ]
    },
    "TCNL": {
        "title": "Tính năng nhập khảo sát hài lòng, an toàn lao động, đình công và quản lý nhân sự.",
        "scope": "Các tính năng nhập điểm ma trận hài lòng nhân viên, tỷ lệ lãnh đạo, báo cáo tai nạn lao động, nhật ký đình công (không hiển thị chỉ tiêu cơ cấu nhân sự theo lĩnh vực CT 4.47).",
        "requirements": [
            ['YCQ-TCNL-WEB-01', 'Rút chỉ tiêu nhân sự theo lĩnh vực (CT 4.47).', 'Cao', 'Không hiển thị lát cắt nhân sự theo lĩnh vực công việc (Pilots/Cabin/Ground) trên Dashboard quy mô nhân sự (GRI 2-7). Chỉ hiển thị nhân sự theo giới tính, độ tuổi, khu vực và hợp đồng.'],
            ['YCQ-TCNL-WEB-02', 'Ma trận điểm hài lòng nhân viên (GRI 405-1).', 'Cao', 'Thiết lập bảng nhập điểm hài lòng gồm 8 tiêu chí: Hài lòng chung, Đặc điểm công việc, Điều kiện làm việc, Đào tạo & Thăng tiến, Ban lãnh đạo, Quan hệ đồng nghiệp, Thu nhập, Phúc lợi. Hỗ trợ nhập theo từng ban/đơn vị hoặc nhập Toàn tổng công ty.'],
            ['YCQ-TCNL-WEB-03', 'Khai báo tai nạn lao động (GRI 403-9) và đình công (GRI 2-30).', 'Cao', 'Đình công: Nhập thông tin Có/Không xảy ra đình công (hasStrike), Số lượng vụ đình công (strikeCount), và Mô tả chi tiết vụ việc (strikeDescription). An toàn lao động: Nút tải lên chứng từ/tệp báo cáo tai nạn lao động định kỳ.']
        ],
        "traceability": [
            ['YCQ-TCNL-WEB-01', 'Hình 01: Dashboard Nhân lực', 'Biểu đồ quy mô nhân sự', 'Xác nhận không còn phân tích theo lĩnh vực công việc.'],
            ['YCQ-TCNL-WEB-02', 'Hình 02: Màn hình Điểm hài lòng', 'Ma trận điểm khảo sát', 'Kiểm thử nhập điểm hài lòng cho 8 tiêu chí và tính toán giá trị trung bình.'],
            ['YCQ-TCNL-WEB-03', 'Hình 03: Màn hình Tai nạn & Đình công', 'Số liệu an toàn & tệp đính kèm', 'Kiểm thử khai báo sự việc đình công và upload file báo cáo tai nạn.']
        ],
        "uat": [
            ['YCQ-TCNL-WEB-01', 'Kiểm thử ẩn chỉ tiêu CT 4.47', 'Kiểm tra biểu đồ quy mô nhân sự để đảm bảo không hiển thị phân tích theo lĩnh vực Pilots/Cabin/Ground.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-TCNL-WEB-02', 'Kiểm thử nhập điểm hài lòng', 'Nhập ma trận điểm khảo sát cho Ban TCNL, kiểm tra lưu trữ và đồng bộ dữ liệu.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-TCNL-WEB-03', 'Kiểm thử khai báo đình công', 'Khai báo có xảy ra 1 vụ đình công tự phát, kiểm tra hiển thị. Thử tải tệp báo cáo tai nạn lao động.', 'Đạt', '—', 'Hoàn thành']
        ]
    },
    "Kỹ thuật": {
        "title": "Tính năng cấu hình định mức phát thải, nhập dữ liệu SAF và import Excel ReFuelEU.",
        "scope": "Các tính năng cấu hình định mức phát thải (ETS/CORSIA), nhập chi phí SAF và import tệp Excel ReFuelEU Aviation 16 cột.",
        "requirements": [
            ['YCQ-TKT-WEB-01', 'Thông số định mức & phát thải CO2.', 'Cao', 'Form nhập liệu các tham số cấu hình: Tỷ lệ SAF quy định (safPlannedRatio), Hạn ngạch ETS (quotaEts), Baseline CORSIA (baselineCorsia), Đơn giá tín chỉ CO2 vượt mức (priceEua), Phụ phí mua SAF (surchargeSaf).'],
            ['YCQ-TKT-WEB-02', 'Import Excel ReFuelEU chuẩn 16 cột.', 'Cao', 'Hỗ trợ import tệp Excel chuẩn 16 cột thông tin gồm: Mã sân bay ICAO, Nhà cung cấp, MST nhà cung cấp, Số lô (Batch), Lượng mua (Tấn), Phân loại SAF, Quy trình, Nguyên liệu, Nguồn gốc, Hệ số phát thải vòng đời, và 6 cột lượng kê khai bù trừ dưới các MBMs (EU, CH, UK, CORSIA, MBM1, MBM2).']
        ],
        "traceability": [
            ['YCQ-TKT-WEB-01', 'Hình 01: Màn hình Định mức', 'Tham số cấu hình phát thải', 'Kiểm thử nhập và lưu các chỉ số định mức và phụ phí.'],
            ['YCQ-TKT-WEB-02', 'Hình 02: Màn hình nạp SAF chi tiết', 'Bảng Excel 16 cột import', 'Kiểm thử kéo thả file Excel ReFuelEU và xác nhận tính hợp lệ của dữ liệu.']
        ],
        "uat": [
            ['YCQ-TKT-WEB-01', 'Kiểm thử cấu hình định mức', 'Nhập tỷ lệ SAF kế hoạch 2.5%, hạn ngạch 16200 tấn, kiểm tra tính toán chi phí đền bù CO2 tương ứng.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-TKT-WEB-02', 'Kiểm thử import Excel ReFuelEU', 'Chọn file Excel mẫu chứa thông tin lô SAF, kiểm tra hệ thống đọc đúng 16 cột dữ liệu và lưu thành công.', 'Đạt', '—', 'Hoàn thành']
        ]
    },
    "Bông Sen Vàng": {
        "title": "Tính năng nhập liệu hội viên mới và Dashboard tăng trưởng Bông Sen Vàng.",
        "scope": "Các tính năng nhập liệu số lượng hội viên mới kế hoạch và thực tế của Trung tâm Bông Sen Vàng (chỉ tiêu Airline B-2).",
        "requirements": [
            ['YCQ-BSV-WEB-01', 'Nhập liệu chỉ tiêu Hội viên mới (Airline B-2).', 'Cao', 'Form nhập liệu gồm 3 trường dữ liệu: Số lượng hội viên mới kế hoạch (npsTarget), Điểm/Số thực tế đạt được (npsActual), và Số lượng khảo sát thu về (customerSurveys).'],
            ['YCQ-BSV-WEB-02', 'Dashboard tăng trưởng hội viên.', 'Cao', 'Hiển thị biểu đồ tăng trưởng số lượng hội viên mới theo tháng trong năm, phân tách theo hạng thẻ hội viên. Dữ liệu lấy tự động từ hệ thống CLM.']
        ],
        "traceability": [
            ['YCQ-BSV-WEB-01', 'Hình 01: Màn hình Nhập liệu hội viên', 'Số liệu hội viên Bông Sen Vàng', 'Kiểm thử nhập và lưu số lượng hội viên mới kế hoạch/thực tế.'],
            ['YCQ-BSV-WEB-02', 'Hình 02: Dashboard KPI Airline B-2', 'Biểu đồ Metabase', 'Kiểm thử hiển thị biểu đồ xu hướng hội viên theo tháng.']
        ],
        "uat": [
            ['YCQ-BSV-WEB-01', 'Kiểm thử nhập số liệu hội viên', 'Nhập các số liệu mục tiêu và thực tế hội viên mới, kiểm tra dữ liệu lưu thành công.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-BSV-WEB-02', 'Kiểm thử Dashboard hội viên', 'Kiểm tra hiển thị biểu đồ xu hướng tăng trưởng hội viên theo các hạng thẻ khác nhau.', 'Đạt', '—', 'Hoàn thành']
        ]
    },
    "An Toàn Chất Lượng": {
        "title": "Tính năng nhập tiếng ồn sân bay, phát thải nội địa và nhật ký sự cố an toàn.",
        "scope": "Các tính năng nhập mức ồn sân bay (Airline E-1), lượng phát thải Scope 1 Domestic và nhật ký sự cố an toàn lao động (GRI 403-2) của Ban ATCL.",
        "requirements": [
            ['YCQ-ATCL-WEB-01', 'Nhập liệu chỉ số tiếng ồn & phát thải.', 'Cao', 'Form nhập liệu gồm: Chỉ số tiếng ồn sân bay (noiseLevel) lấy từ AMOS, Lượng phát thải Scope 1 Domestic (ghgEmissions) lấy từ FIMS/Excel.'],
            ['YCQ-ATCL-WEB-02', 'Nhật ký sự cố an toàn (GRI 403-2).', 'Cao', 'Giao diện bảng cho phép thêm mới, sửa đổi danh sách sự cố an toàn lao động/dịch vụ gồm: Mã sự cố, Ngày xảy ra, Tiêu đề sự cố, Trạng thái xử lý (Open | Investigating | Closed).']
        ],
        "traceability": [
            ['YCQ-ATCL-WEB-01', 'Hình 01: Màn hình Nhập liệu ATCL', 'Số liệu mức ồn và khí thải', 'Kiểm thử nhập chỉ số ồn sân bay và lượng phát thải Scope 1.'],
            ['YCQ-ATCL-WEB-02', 'Hình 02: Màn hình Nhật ký sự cố an toàn', 'Danh sách sự cố an toàn', 'Kiểm thử nhập thông tin sự cố an toàn lao động và đổi trạng thái xử lý.']
        ],
        "uat": [
            ['YCQ-ATCL-WEB-01', 'Kiểm thử nhập mức ồn và phát thải', 'Nhập mức ồn 85dB và lượng khí thải 12500 tấn CO2, kiểm tra ghi nhận dữ liệu.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-ATCL-WEB-02', 'Kiểm thử nhật ký sự cố', 'Nhập một vụ tai nạn lao động tại sân đỗ, kiểm tra lưu trữ và theo dõi biểu đồ sự cố.', 'Đạt', '—', 'Hoàn thành']
        ]
    },
    "TCKT-KHPT": {
        "title": "Tính năng Dashboard tài chính ESG và kết xuất báo cáo kiểm toán tài chính.",
        "scope": "Các tính năng hiển thị Dashboard phân tích chi phí carbon, phụ phí SAF và kết xuất dữ liệu tài chính ESG ra Excel của Ban TCKT-KHPT.",
        "requirements": [
            ['YCQ-TCKT-WEB-01', 'Dashboard tài chính ESG.', 'Cao', 'Hiển thị biểu đồ phân tích chi phí mua tín chỉ carbon (EU ETS / CORSIA), phụ phí SAF và dòng tiền đầu tư cho công nghệ tiết kiệm năng lượng. Cho phép lọc theo năm tài chính.'],
            ['YCQ-TCKT-WEB-02', 'Xem & tải dữ liệu tài chính phục vụ báo cáo.', 'Cao', 'Giao diện cho phép Ban TCKT kết xuất các bảng biểu tài chính liên quan đến phát thải và SAF ra định dạng Excel (.xlsx) phục vụ kiểm toán.']
        ],
        "traceability": [
            ['YCQ-TCKT-WEB-01', 'Hình 01: Dashboard Tài chính ESG', 'Biểu đồ Metabase', 'Kiểm thử hiển thị chi phí carbon và dòng tiền đầu tư phát triển bền vững.'],
            ['YCQ-TCKT-WEB-02', 'Hình 02: Màn hình Kết xuất báo cáo', 'Báo cáo kiểm toán Excel', 'Kiểm thử tải tệp dữ liệu tài chính dạng Excel.']
        ],
        "uat": [
            ['YCQ-TCKT-WEB-01', 'Kiểm thử Dashboard tài chính', 'Kiểm tra biểu đồ chi phí carbon hiển thị đúng số liệu đối soát từ Metabase.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-TCKT-WEB-02', 'Kiểm thử kết xuất báo cáo', 'Nhấn nút tải báo cáo tài chính ESG, kiểm tra tệp Excel tải xuống chứa đầy đủ số liệu và định dạng hợp lệ.', 'Đạt', '—', 'Hoàn thành']
        ]
    }
}

def highlight_paragraph(p):
    # Set text and highlight all runs inside paragraph
    for run in p.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def highlight_cell(cell):
    # Highlight all runs in all paragraphs of a cell
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

docx_files = glob.glob(os.path.join(appendix_dir, "*.docx"))
print(f"Found {len(docx_files)} files to update with highlight.")

for fpath in docx_files:
    basename = os.path.basename(fpath)
    if "OneDrive" in basename:
        continue # Skip zip file
        
    print("-" * 60)
    print(f"Highlighting edits in: {basename}")
    
    dept_key = None
    for k in DEPT_DATA.keys():
        if k in basename:
            dept_key = k
            break
            
    if not dept_key:
        print(f"Warning: No match found for file {basename}. Skipping.")
        continue
        
    data = DEPT_DATA[dept_key]
    doc = docx.Document(fpath)
    
    # 1. Update paragraphs with highlight
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        # A. Update workflow warning (Offline approval rule)
        if "Không có" in text and idx > 0 and "5. Cảnh báo và workflow khi nhập liệu" in doc.paragraphs[idx-1].text:
            p.text = "Quy trình phê duyệt số liệu được thực hiện ngoại tuyến (offline) theo quy chế vận hành thực tế. Hệ thống ghi nhận số liệu trực tiếp vào cơ sở dữ liệu Clean DB và cập nhật tức thời lên Dashboard."
            highlight_paragraph(p)
            
        # B. Update UAT title
        if "Chức năng nhập liệu an toàn chất lượng, Import Excel và Dashboard Xã hội" in text:
            p.text = data["title"]
            highlight_paragraph(p)
            
        # C. Update UAT scope
        if "Các chỉ tiêu tiếng ồn, phát thải nội địa và sự cố an toàn" in text:
            p.text = data["scope"]
            highlight_paragraph(p)

    # 2. Update tables with highlight
    for t_idx, table in enumerate(doc.tables):
        header = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
        
        # A. Target Requirements Table
        if len(header) >= 4 and "Mã yêu cầu" in header[0] and "Tiêu chí chấp nhận" in header[3]:
            while len(table.rows) > 1:
                tr = table.rows[1]._tr
                table._tbl.remove(tr)
            for req in data["requirements"]:
                row_cells = table.add_row().cells
                for col_idx, val in enumerate(req):
                    row_cells[col_idx].text = val
                    highlight_cell(row_cells[col_idx])
            print(f"  Highlighted Requirements table ({len(table.rows)} rows)")
            
        # B. Target Traceability Table
        elif len(header) >= 4 and "Mã yêu cầu" in header[0] and "Test case" in header[3]:
            while len(table.rows) > 1:
                tr = table.rows[1]._tr
                table._tbl.remove(tr)
            for trace in data["traceability"]:
                row_cells = table.add_row().cells
                for col_idx, val in enumerate(trace):
                    row_cells[col_idx].text = val
                    highlight_cell(row_cells[col_idx])
            print(f"  Highlighted Traceability table ({len(table.rows)} rows)")
            
        # C. Target UAT Table
        elif len(header) >= 6 and "Mã yêu cầu" in header[0] and "Trạng thái" in header[5]:
            while len(table.rows) > 1:
                tr = table.rows[1]._tr
                table._tbl.remove(tr)
            for uat in data["uat"]:
                row_cells = table.add_row().cells
                for col_idx, val in enumerate(uat):
                    row_cells[col_idx].text = val
                    highlight_cell(row_cells[col_idx])
            print(f"  Highlighted UAT table ({len(table.rows)} rows)")
            
    doc.save(fpath)
    print(f"Success! Highlighted & saved: {basename}")

print("\nAll docx files successfully highlighted and saved!")

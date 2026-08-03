import docx
import os

doc_path = "/Users/congnguyen/Library/CloudStorage/OneDrive-Personal/Work/04. Company/Synodus/GOV/01.VNA/Document/Phụ lục/Phụ lục_Ban KHPT_ver1.docx"
output_path = "/Users/congnguyen/Library/CloudStorage/OneDrive-Personal/Work/04. Company/Synodus/GOV/01.VNA/VNA/scratch/test_updated_KHPT.docx"

doc = docx.Document(doc_path)

# Let's inspect headings and print out paragraph indices where we perform text replacement
print("=== Paragraph Replacements ===")
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    # Replace Bông Sen Vàng disclaimers or workflow descriptions
    if "Không có" in text and idx > 0 and "5. Cảnh báo và workflow khi nhập liệu" in doc.paragraphs[idx-1].text:
        print(f"P {idx} (workflow): {text} -> Replacing with Offline workflow disclaimer")
        p.text = "Quy trình phê duyệt số liệu được thực hiện ngoại tuyến (offline) theo quy chế vận hành thực tế. Hệ thống ghi nhận số liệu trực tiếp vào cơ sở dữ liệu Clean DB và cập nhật tức thời lên Dashboard."
    if "Chức năng nhập liệu an toàn chất lượng, Import Excel và Dashboard Xã hội" in text:
        print(f"P {idx} (UAT Title): {text} -> Replacing with KHPT business description")
        p.text = "Tính năng Dashboard tổng quan, Quản lý mục tiêu KPI và Kho tài liệu PTBV chung."
    if "Các chỉ tiêu tiếng ồn, phát thải nội địa và sự cố an toàn" in text:
        print(f"P {idx} (UAT Scope): {text[:50]}... -> Replacing with KHPT UAT scope")
        p.text = "Các tính năng hiển thị Dashboard tổng quan chung, 3 trụ cột E, S, G, Quản lý KPI và Kho tài liệu PTBV chung của Ban KHPT."

print("\n=== Table Replacements ===")
for t_idx, table in enumerate(doc.tables):
    header = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
    
    # 1. Target Requirements Table
    if len(header) >= 4 and "Mã yêu cầu" in header[0] and "Tiêu chí chấp nhận" in header[3]:
        print(f"Table {t_idx}: Requirements list detected. Rows count: {len(table.rows)}")
        # Delete placeholder rows after header (row 0)
        while len(table.rows) > 1:
            # Delete row 1 (which is the first data row)
            # In python-docx, table._tbl.remove(row._tr) is used to delete rows
            tr = table.rows[1]._tr
            table._tbl.remove(tr)
            
        # Add new data rows
        requirements = [
            ['YCQ-KHPT-WEB-01', 'Giao diện Dashboard điều hành tổng quan chung (Executive Dashboard).', 'Cao', 
             'Hiển thị biểu đồ phân tích tích hợp (Metabase iframe) gồm 4 tab chuyển đổi: Tổng quan chung (ALL), Environmental (E), Social (S), và Governance (G). Hỗ trợ nút xem chi tiết trên Metabase và bảng danh sách chỉ tiêu chi tiết bên dưới.'],
            ['YCQ-KHPT-WEB-02', 'Quy trình Đề xuất & Thiết lập mục tiêu KPI.', 'Cao', 
             'Giao diện hỗ trợ quản lý danh mục KPI có các trường: Mã chỉ tiêu, Tên chỉ tiêu, Đơn vị tính, Số kế hoạch, Số thực hiện, Tiến độ (%), Đánh giá (Đạt/Không đạt), và Đơn vị chủ trì. Cho phép KHPT gửi yêu cầu làm kế hoạch KPI năm xuống các phòng ban, theo dõi tiến độ nộp và phê duyệt/yêu cầu nộp lại số liệu kế hoạch trực tiếp trên hệ thống (Kế hoạch lưu vào mục riêng cho KHPT quản lý).'],
            ['YCQ-KHPT-WEB-03', 'Quản trị Kho tài liệu Phát triển bền vững chung.', 'Trung bình', 
             'Tiếp nhận yêu cầu upload tài liệu của các phòng ban dưới dạng tệp đính kèm. Hỗ trợ KHPT duyệt lưu trữ (vào thư mục dùng chung) hoặc từ chối chỉnh sửa.']
        ]
        for req in requirements:
            row_cells = table.add_row().cells
            for col_idx, val in enumerate(req):
                row_cells[col_idx].text = val
        print(f"Table {t_idx} requirements updated. New rows: {len(table.rows)}")

    # 2. Target Design/Traceability Table
    elif len(header) >= 4 and "Mã yêu cầu" in header[0] and "Test case" in header[3]:
        print(f"Table {t_idx}: Traceability detected. Rows count: {len(table.rows)}")
        while len(table.rows) > 1:
            tr = table.rows[1]._tr
            table._tbl.remove(tr)
            
        traceability = [
            ['YCQ-KHPT-WEB-01', 'Hình 01: Giao diện dashboard tổng quan', 'Dữ liệu Metabase nhúng', 'Kiểm thử hiển thị biểu đồ và bảng danh sách chỉ tiêu.'],
            ['YCQ-KHPT-WEB-02', 'Hình 02: Màn hình quản lý KPI', 'Dữ liệu chỉ tiêu KPI', 'Kiểm thử cập nhật số kế hoạch và theo dõi tiến độ.'],
            ['YCQ-KHPT-WEB-03', 'Hình 03: Màn hình kho tài liệu', 'Tệp đính kèm pdf/docx', 'Kiểm thử tải lên tài liệu và nút kiểm duyệt lưu trữ.']
        ]
        for trace in traceability:
            row_cells = table.add_row().cells
            for col_idx, val in enumerate(trace):
                row_cells[col_idx].text = val
        print(f"Table {t_idx} traceability updated. New rows: {len(table.rows)}")

    # 3. Target UAT Table
    elif len(header) >= 6 and "Mã yêu cầu" in header[0] and "Trạng thái" in header[5]:
        print(f"Table {t_idx}: UAT table detected. Rows count: {len(table.rows)}")
        while len(table.rows) > 1:
            tr = table.rows[1]._tr
            table._tbl.remove(tr)
            
        uat_cases = [
            ['YCQ-KHPT-WEB-01', 'Kiểm thử Dashboard tổng quan', 'Mở trang Tổng quan điều hành, kiểm tra hiển thị iframe Metabase và các tab E, S, G. Kiểm tra bảng chỉ tiêu bên dưới.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-KHPT-WEB-02', 'Kiểm thử quản lý KPI', 'Mở chức năng Quản lý KPI, kiểm tra nhập số kế hoạch KPI cho các chỉ tiêu.', 'Đạt', '—', 'Hoàn thành'],
            ['YCQ-KHPT-WEB-03', 'Kiểm thử Kho tài liệu', 'Mở chức năng Kho tài liệu PTBV chung, kiểm tra tải lên tài liệu và kiểm duyệt lưu trữ.', 'Đạt', '—', 'Hoàn thành']
        ]
        for case in uat_cases:
            row_cells = table.add_row().cells
            for col_idx, val in enumerate(case):
                row_cells[col_idx].text = val
        print(f"Table {t_idx} UAT updated. New rows: {len(table.rows)}")

doc.save(output_path)
print("Saved tested updated document to:", output_path)

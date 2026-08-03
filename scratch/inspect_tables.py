import docx
import os

doc_path = "/Users/congnguyen/Library/CloudStorage/OneDrive-Personal/Work/04. Company/Synodus/GOV/01.VNA/Document/Phụ lục/Phụ lục_Ban KHPT_ver1.docx"

if not os.path.exists(doc_path):
    print("Doc file not found at:", doc_path)
    exit(1)

doc = docx.Document(doc_path)
print(f"Document: {os.path.basename(doc_path)}")
print(f"Number of paragraphs: {len(doc.paragraphs)}")
print(f"Number of tables: {len(doc.tables)}")

# Print some paragraphs at the end of the document
print("\n--- Last 30 paragraphs: ---")
for idx in range(max(0, len(doc.paragraphs)-30), len(doc.paragraphs)):
    p = doc.paragraphs[idx]
    if p.text.strip():
        print(f"P {idx}: {p.text.strip()}")

# Inspect tables
print("\n--- Tables Inspection: ---")
for t_idx, table in enumerate(doc.tables):
    print(f"Table {t_idx} (rows: {len(table.rows)}, cols: {len(table.columns)})")
    # Print first row of the table
    first_row_text = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
    print(f"  Header: {first_row_text[:4]}")
    # Check if this table has BSV placeholder
    has_bsv = False
    for row in table.rows:
        row_text = "".join(cell.text for cell in row.cells)
        if "Bông Sen Vàng" in row_text or "BSV" in row_text or "Airline B-2" in row_text:
            has_bsv = True
            break
    if has_bsv:
        print(f"  -> TABLE {t_idx} HAS BSV PLACEHOLDER")
        # Print some rows
        for r_idx, row in enumerate(table.rows[:3]):
            cell_texts = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            print(f"    Row {r_idx}: {cell_texts[:4]}")
